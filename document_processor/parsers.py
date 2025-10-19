import os
import unicodedata
from datetime import datetime
from typing import Dict, List, Tuple
from pypdf import PdfReader
from docx import Document


class DocumentParser:
    """Handles parsing of different document types"""

    @staticmethod
    def parse_pdf(file_path: str) -> Tuple[str, Dict]:
        """Parse PDF file and extract text and metadata"""
        try:
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"

            metadata = {
                'num_pages': len(reader.pages),
                'title': reader.metadata.get('/Title', '') if reader.metadata else '',
                'author': reader.metadata.get('/Author', '') if reader.metadata else '',
                'creation_date': str(reader.metadata.get('/CreationDate', '')) if reader.metadata else '',
            }

            return text, metadata
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")

    @staticmethod
    def parse_docx(file_path: str) -> Tuple[str, Dict]:
        """Parse DOCX file and extract text and metadata"""
        try:
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

            metadata = {
                'num_paragraphs': len(doc.paragraphs),
                'author': doc.core_properties.author or '',
                'title': doc.core_properties.title or '',
                'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                'modified': str(doc.core_properties.modified) if doc.core_properties.modified else '',
            }

            return text, metadata
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")

    @staticmethod
    def extract_metadata(file_path: str, file_type: str) -> Dict:
        """Extract basic file metadata"""
        stat_info = os.stat(file_path)
        return {
            'file_name': os.path.basename(file_path),
            'file_size': stat_info.st_size,
            'file_type': file_type,
            'created_at': datetime.fromtimestamp(stat_info.st_ctime).isoformat(),
            'modified_at': datetime.fromtimestamp(stat_info.st_mtime).isoformat(),
        }


class TextNormalizer:
    """Handles text normalization and cleaning"""

    @staticmethod
    def normalize_text(text: str) -> str:
        """Normalize text by removing accents and special characters"""
        text = unicodedata.normalize('NFKD', text)
        text = ''.join([c for c in text if not unicodedata.combining(c)])
        return text

    @staticmethod
    def clean_text(text: str) -> str:
        """Clean text by removing extra whitespace and normalizing"""
        text = ' '.join(text.split())
        text = text.strip()
        return text

    @staticmethod
    def process_text(text: str, normalize: bool = True) -> str:
        """Process text with cleaning and optional normalization"""
        text = TextNormalizer.clean_text(text)
        if normalize:
            text = TextNormalizer.normalize_text(text)
        return text


class TextChunker:
    """Handles text chunking with configurable overlap"""

    @staticmethod
    def chunk_text(text: str, chunk_size: int, overlap: int) -> List[Dict]:
        """Split text into overlapping chunks"""
        chunks = []
        text_length = len(text)
        start = 0
        chunk_index = 0

        while start < text_length:
            end = start + chunk_size

            if end >= text_length:
                chunk_content = text[start:]
            else:
                chunk_content = text[start:end]
                last_space = chunk_content.rfind(' ')
                if last_space != -1 and last_space > chunk_size * 0.5:
                    end = start + last_space
                    chunk_content = text[start:end]

            chunks.append({
                'content': chunk_content.strip(),
                'chunk_index': chunk_index,
                'start_position': start,
                'end_position': end,
            })

            start = end - overlap
            chunk_index += 1

        return chunks
